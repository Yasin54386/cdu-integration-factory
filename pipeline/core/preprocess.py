"""Supporting-file preprocessing for prompt assembly (spec §6).

Copilot CLI consumes text. Each file type is reduced to text:
  .sql               → verbatim, full content
  .docx              → text via python-docx (paragraphs + tables)
  .pdf               → text via pypdf; empty extraction = FAIL (scanned PDF)
  .xlsx              → markdown table via openpyxl
  .csv/.dat/.txt     → first 20 rows + truncation note with total row count

Per-file and whole-prompt character caps FAIL loudly rather than silently
truncating business rules.
"""

from __future__ import annotations

from pathlib import Path

PER_FILE_CHAR_CAP = 8_000
PROMPT_CHAR_CAP = 48_000
SAMPLE_ROW_LIMIT = 20


class PreprocessError(ValueError):
    """Raised when a file cannot be reduced to usable prompt text."""


def _preprocess_sql(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _preprocess_docx(path: Path) -> str:
    import docx  # python-docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _preprocess_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    text = "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    if not text:
        raise PreprocessError(
            f"{path.name}: no extractable text — scanned PDF unsupported in v1; "
            "provide a text-based PDF or .docx instead"
        )
    return text


def _preprocess_xlsx(path: Path) -> str:
    import openpyxl

    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        rows = [
            ["" if cell is None else str(cell) for cell in row]
            for row in sheet.iter_rows(values_only=True)
        ]
        if not rows:
            continue
        lines.append(f"### Sheet: {sheet.title}")
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("|" + "---|" * len(rows[0]))
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")
    workbook.close()
    return "\n".join(lines)


def _preprocess_sample(path: Path) -> str:
    lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    total = len(lines)
    head = lines[:SAMPLE_ROW_LIMIT]
    text = "\n".join(head)
    if total > SAMPLE_ROW_LIMIT:
        text += (
            f"\n[truncated: showing first {SAMPLE_ROW_LIMIT} of {total} rows]"
        )
    return text


_HANDLERS = {
    ".sql": _preprocess_sql,
    ".docx": _preprocess_docx,
    ".pdf": _preprocess_pdf,
    ".xlsx": _preprocess_xlsx,
    ".csv": _preprocess_sample,
    ".dat": _preprocess_sample,
    ".txt": _preprocess_sample,
}


def preprocess_file(path: Path) -> str:
    """Reduce one supporting file to prompt text, enforcing the per-file cap."""
    handler = _HANDLERS.get(path.suffix.lower())
    if handler is None:
        raise PreprocessError(
            f"{path.name}: unsupported file type '{path.suffix}' "
            f"(supported: {sorted(_HANDLERS)})"
        )
    text = handler(path)
    if len(text) > PER_FILE_CHAR_CAP:
        raise PreprocessError(
            f"{path.name}: preprocessed content is {len(text)} chars, "
            f"over the {PER_FILE_CHAR_CAP}-char per-file cap — trim the file or "
            "split it rather than relying on silent truncation"
        )
    return text


def enforce_prompt_cap(prompt: str) -> str:
    if len(prompt) > PROMPT_CHAR_CAP:
        raise PreprocessError(
            f"assembled prompt is {len(prompt)} chars, over the "
            f"{PROMPT_CHAR_CAP}-char cap — trim supporting files or intent notes"
        )
    return prompt
